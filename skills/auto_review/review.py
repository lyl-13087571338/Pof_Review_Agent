import fitz  # PyMuPDF
from docx import Document
import os
import yaml
import requests
from pathlib import Path

# Load configuration
CONFIG_PATH = Path(__file__).resolve().parent.parent.parent / "config" / "config.yaml"
with open(CONFIG_PATH, "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

llm_config = config.get("llm", {})

LLM_API_KEY = llm_config.get("api_key")
LLM_API_BASE = llm_config.get("api_base", "https://open.bigmodel.cn/api/paas/v4")
LLM_MODEL = llm_config.get("model", "glm-3-turbo")
OUTPUT_FOLDER = config["paths"]["review_folder"]
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def pdf_to_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def generate_review(text, title):
    prompt = f"你是审稿专家，请基于以下稿件生成 POF 审稿意见：\n{text}\n要求输出 Word 可直接展示的格式，包括优点、缺点、改进建议。"

    api_base = LLM_API_BASE.rstrip("/")
    url = f"{api_base}/chat/completions"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LLM_API_KEY}",
    }
    payload = {
        "model": LLM_MODEL,
        "messages": [{"role": "user", "content": prompt}],
    }

    print(f"[INFO] Calling Zhipu API with model: {LLM_MODEL}")
    response = requests.post(url, json=payload, headers=headers, timeout=300)
    
    if response.status_code != 200:
        print(f"[ERROR] API Error {response.status_code}: {response.text}")
    
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"]


def save_to_word(review_text, title):
    word_path = os.path.join(OUTPUT_FOLDER, f"{title}_审稿意见.docx")
    doc = Document()
    doc.add_heading("审稿意见", 0)
    doc.add_paragraph(review_text)
    doc.save(word_path)
    return word_path


def review_pdf(pdf_path):
    text = pdf_to_text(pdf_path)
    title = os.path.splitext(os.path.basename(pdf_path))[0]
    review_text = generate_review(text, title)
    word_file = save_to_word(review_text, title)
    return word_file


if __name__ == "__main__":
    import sys
    pdf_file = sys.argv[1]
    word_file = review_pdf(pdf_file)
    print("Review saved at:", word_file)
